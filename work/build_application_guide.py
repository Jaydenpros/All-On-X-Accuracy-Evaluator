import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "OnXTrue_STL_Viewer_Application_Guide.docx"
SCREENSHOT = Path(os.environ["ONXTRUE_GUIDE_SCREENSHOT"]) if os.environ.get("ONXTRUE_GUIDE_SCREENSHOT") else None

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "17324D"
GREEN = "2E9D6B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F4EE"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
DARK = "1F2933"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=False, color=DARK, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_font(run, size=9, color=MID_GRAY)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    return paragraph


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_font(first, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest)
    else:
        run = p.add_run(text)
        set_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_font(p.add_run(text))
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}  ")
    set_font(label_run, size=10, bold=True, color=accent)
    text_run = p.add_run(text)
    set_font(text_run, size=10.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    header = table.rows[0].cells
    header[0].text = "Area"
    header[1].text = "Implementation"
    set_repeat_table_header(table.rows[0])
    for cell in header:
        set_cell_shading(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10, bold=True, color=WHITE)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            set_font(run, size=10, bold=True, color=DARK_BLUE)
        for run in cells[1].paragraphs[0].runs:
            set_font(run, size=10, color=DARK)
    set_table_geometry(table, [2700, 6660])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "ONXTRUE  |  TECHNICAL APPLICATION GUIDE"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        set_font(run, size=8, bold=True, color=MID_GRAY)
    add_page_number(section.footer.paragraphs[0])

    # Editorial cover
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(105)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("TECHNICAL APPLICATION GUIDE"), size=10, bold=True, color=GREEN)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("OnXTrue STL Viewer"), size=30, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(26)
    set_font(subtitle.add_run("Multi-file dental scan visualization and scanbody isolation"), size=14, color=DARK_BLUE)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(26)
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GREEN)
    borders.append(bottom)
    p_pr.append(borders)
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(date_p.add_run("Version 1.0  •  June 18, 2026"), size=11, bold=True, color=MID_GRAY)
    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(audience.add_run("Prepared as an operator, developer, and deployment reference"), size=9.5, italic=True, color=MID_GRAY)

    doc.add_page_break()

    add_heading(doc, "1. Application overview", 1)
    add_body(doc, "OnXTrue is a browser-based STL visualization application with a Python Flask backend. It is designed for dental scan workflows that compare one Scan Body Library model with one or more scan files in a shared native coordinate system.")
    add_callout(
        doc,
        "CORE WORKFLOW",
        "Load one blue library STL, load multiple green scan STLs, inspect their native alignment, optionally isolate disconnected scanbodies, and control each 3D object independently.",
        LIGHT_GREEN,
        GREEN,
    )
    add_heading(doc, "Key capabilities", 2)
    for item in [
        "Single-file Scan Body Library upload with automatic replacement of the previous library model.",
        "Multiple scan-file upload with green rendering and independent object visibility.",
        "Interactive orbit, pan, zoom, fit-view, reset-view, wireframe, and grid controls.",
        "Native STL coordinate preservation with the library origin used as scene origin.",
        "Z-up coordinate convention: Z is vertical and the grid lies on the XY plane at Z = 0.",
        "Connected-component scanbody isolation with a 5 mm proximity merge rule.",
        "ASCII and binary STL support, including vendor-specific 3Shape binary STL headers.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. Interface overview", 1)
    if SCREENSHOT and SCREENSHOT.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(SCREENSHOT), width=Inches(6.2))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)
        set_font(caption.add_run("Figure 1. OnXTrue scene with the blue library model, green combined scan, Z-up axes, and 3D Objects panel."), size=9, italic=True, color=MID_GRAY)

    add_key_value_table(doc, [
        ("Scan Body Library", "Accepts one STL. The model is blue and defines the coordinate origin."),
        ("Scans", "Accepts multiple STL files. Models are green and retain their original coordinates."),
        ("3D Objects panel", "Lists all loaded or isolated objects with independent visibility toggles."),
        ("Viewer toolbar", "Fit/reset camera, isolate scanbodies, toggle wireframe, and toggle grid."),
        ("Scene statistics", "Shows visible triangle count, scene dimensions, and library/scan object counts."),
    ])

    add_heading(doc, "3. Operating instructions", 1)
    add_heading(doc, "Start the application", 2)
    for step in [
        "Open Terminal and change to the folder where you cloned or copied the project.",
        "Activate the environment with: source .venv/bin/activate",
        "Start Flask with: python app.py",
        "Keep the Terminal window open and browse to http://127.0.0.1:5000/",
    ]:
        add_number(doc, step)
    add_callout(doc, "CODEX PREVIEW", "The Codex in-app browser may use the separately managed preview at http://127.0.0.1:5001/.", LIGHT_BLUE, BLUE)

    add_heading(doc, "Load and inspect models", 2)
    for step in [
        "Upload the reference STL in Scan Body Library. A later upload replaces the current library file.",
        "Upload one or more STL files in Scans. Each original scan appears as a green 3D object.",
        "Use Fit view to frame all visible geometry.",
        "Orbit with left-drag, zoom with the mouse wheel, and pan with right-drag.",
        "Use the 3D Objects switches to compare selected models without deleting them.",
    ]:
        add_number(doc, step)

    doc.add_page_break()
    add_heading(doc, "4. Coordinate and visualization conventions", 1)
    add_key_value_table(doc, [
        ("Scene origin", "The native (0, 0, 0) origin of the Scan Body Library STL."),
        ("Coordinate preservation", "Models are not centered, translated, or arranged side-by-side."),
        ("Z axis", "Vertical axis."),
        ("X and Y axes", "Horizontal plane axes."),
        ("Grid", "XY plane positioned at Z = 0."),
        ("Axis visibility", "Axis helper is scaled to remain highly visible relative to the library model."),
        ("Color convention", "Library = blue; scans and isolated scanbodies = green."),
    ])
    add_callout(
        doc,
        "IMPORTANT",
        "Correct spatial comparison depends on the library and scan STL files sharing the same coordinate frame before upload. OnXTrue preserves coordinates; it does not perform automatic registration.",
        "FFF3D6",
        "8A6500",
    )

    add_heading(doc, "5. Isolate Scanbodies in Scans", 1)
    add_body(doc, "The Isolate Scanbodies command converts each combined scan STL into separately controlled scanbody objects. The original combined scan object is replaced by the isolated results.")
    add_heading(doc, "Processing logic", 2)
    for step in [
        "Read every triangle from the scan geometry.",
        "Build connected components by joining triangles that share vertex coordinates.",
        "Measure proximity between disconnected components using bounding-box filtering and vertex-to-vertex distance checks.",
        "Merge components when their surfaces are within 5 mm of one another.",
        "Sort the final groups by position for repeatable numbering.",
        "Create independent objects named Original Name SB1.stl, Original Name SB2.stl, and so on.",
    ]:
        add_number(doc, step)
    add_callout(
        doc,
        "EXAMPLE RESULT",
        "The provided Test scan.stl contains seven disconnected mesh components. Components 1 and 2 are within 5 mm, so the final result is six scanbody objects.",
        LIGHT_GREEN,
        GREEN,
    )
    add_heading(doc, "Visibility after isolation", 2)
    add_body(doc, "Every isolated scanbody appears in both the Scans file list and the 3D Objects panel. Its visibility switch is synchronized across both locations, and Show All / Hide All affects the complete scene.")

    add_heading(doc, "6. Technical architecture", 1)
    add_key_value_table(doc, [
        ("Backend", "Python with Flask 3.0.3."),
        ("Upload API", "POST /api/upload accepts multipart STL files and returns stored-file metadata."),
        ("File validation", "Checks extension plus binary STL byte structure or ASCII STL signatures."),
        ("File storage", "Uploaded files are stored under the local uploads directory with UUID-prefixed names."),
        ("Frontend", "HTML, CSS, and ES modules."),
        ("3D engine", "Three.js 0.180.0 with WebGLRenderer, OrbitControls, and STLLoader."),
        ("Custom STL handling", "Strict binary STL parser is attempted first to support vendor-specific headers."),
        ("Rendering", "Geometry and interaction are processed locally in the browser after upload."),
    ])

    add_heading(doc, "Request and rendering flow", 2)
    flow = doc.add_table(rows=1, cols=3)
    set_table_geometry(flow, [2800, 2800, 3760])
    labels = [
        ("1  UPLOAD", "Browser sends one or more STL files to Flask."),
        ("2  VALIDATE", "Backend validates and stores accepted files."),
        ("3  RENDER", "Browser parses geometry and creates interactive Three.js meshes."),
    ]
    for cell, (label, detail) in zip(flow.rows[0].cells, labels):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(label), size=9, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p2.add_run(detail), size=9.5, color=DARK)

    doc.add_page_break()
    add_heading(doc, "7. Controls reference", 1)
    add_key_value_table(doc, [
        ("Fit view", "Frames all currently visible models."),
        ("Reset view", "Returns the camera to the default Z-up perspective."),
        ("Isolate Scanbodies", "Separates non-isolated scan objects and merges components within 5 mm."),
        ("Wireframe", "Toggles wireframe rendering for every loaded object."),
        ("Grid", "Shows or hides the XY reference grid."),
        ("Object visibility", "Shows or hides one model without removing it."),
        ("Show All / Hide All", "Changes visibility for the entire 3D Objects list."),
        ("Clear scans", "Removes all scan and isolated scanbody objects while preserving the library."),
        ("Remove object", "Deletes one object from the current browser scene."),
    ])

    add_heading(doc, "8. Constraints and limitations", 1)
    for item in [
        "The application is a local development deployment and uses Flask’s development server.",
        "Three.js is currently loaded from jsDelivr, so internet access is required when the page first loads.",
        "Uploads are limited to 500 MB per request.",
        "Uploaded files persist in the local uploads directory until manually cleaned.",
        "Scanbody isolation is geometry-based; the 5 mm threshold is fixed in the current interface.",
        "The isolation proximity test uses mesh vertices, not exact triangle-to-triangle surface distance.",
        "OnXTrue does not automatically align, register, or transform scans against the library.",
        "Model state is browser-session based and is not restored after a page reload.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "9. Troubleshooting", 1)
    add_key_value_table(doc, [
        ("Address already in use", "Another process is using the port. Stop that server with Ctrl+C or use another port."),
        ("Codex cannot reach port 5000", "Use the Codex-managed preview on port 5001 while Chrome continues on port 5000."),
        ("STL cannot be rendered", "Confirm it is a valid ASCII or binary STL. The custom binary parser supports 3Shape-style headers."),
        ("Objects appear misaligned", "Verify the files were exported in the same coordinate frame; OnXTrue does not register models."),
        ("Isolation button disabled", "Load at least one non-isolated scan STL."),
        ("Old interface remains visible", "Reload the page; versioned JavaScript and CSS URLs are used to avoid stale browser caching."),
    ])

    add_heading(doc, "10. Recommended future enhancements", 1)
    for item in [
        "User-adjustable scanbody proximity threshold.",
        "Export isolated scanbodies as individual STL files.",
        "Rigid registration and library-to-scan alignment tools.",
        "Persistent projects with database-backed metadata.",
        "Measurement, clipping-plane, and section-analysis tools.",
        "Offline bundling of Three.js dependencies.",
        "Production WSGI deployment, authentication, and managed file retention.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Appendix A. Project structure", 1)
    add_key_value_table(doc, [
        ("app.py", "Flask application, upload validation, storage, and file-serving routes."),
        ("templates/index.html", "Application layout and controls."),
        ("static/app.js", "Three.js scene, STL parsing, controls, visibility, coordinate handling, and isolation."),
        ("static/styles.css", "Responsive dark user-interface styling."),
        ("requirements.txt", "Python dependency pin for Flask."),
        ("uploads/", "Runtime storage for uploaded STL files."),
    ])

    doc.core_properties.title = "OnXTrue STL Viewer Application Guide"
    doc.core_properties.subject = "Technical and operating documentation"
    doc.core_properties.author = "OnXTrue Project"
    doc.core_properties.keywords = "STL, Three.js, Flask, scanbody, dental scans, 3D viewer"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
